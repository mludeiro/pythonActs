from pathlib import Path
import mimetypes
import pdfplumber
import docx
from abc import ABC, abstractmethod

class Document(ABC):
    def __init__(self, filepath: Path):
        self.filepath = filepath

    @property
    def mime_type(self) -> tuple[str|None, str|None]:
        """Returns a tuple of the form (mimetype, encoding). If either the type or the encoding
        cannot be guessed the value will be None."""
        return mimetypes.guess_type(self.filepath)

    @abstractmethod
    def text(self) -> str:
        pass

    @abstractmethod
    def author(self) -> str:
        pass

class PDFDocument(Document):
    def __init__(self, filepath: Path):
        super().__init__(filepath)
        self.document = pdfplumber.open(self.filepath)

    @property
    def author(self) -> str:
        return self.document.metadata.get("Author", "")

    def text(self) -> str:
        return "\n".join(page.extract_text_simple() for page in self.document.pages)


class DocXDocument(Document):
    def __init__(self, filepath: Path):
        super().__init__(filepath)
        self.document = docx.Document(self.filepath)

    @property
    def author(self) -> str:
        return self.document.core_properties.author

    def text(self) -> str:
        return "\n".join(para.text for para in self.document.paragraphs)
    
if __name__ == "__main__":
    doc = PDFDocument(Path("A.pdf"))
    print(">>>>>>>>>>>>>>>>>>>>>>>>>>>")
    print(doc.author)
    print(">>>>>>>>>>>>>>>>>>>>>>>>>>>")
    print(doc.text())
    print("<<<<<<<<<<<<<<<<<<<<<<<<<<<")

    doc = DocXDocument(Path("B.docx"))
    print(">>>>>>>>>>>>>>>>>>>>>>>>>>>")
    print(doc.author)
    print(">>>>>>>>>>>>>>>>>>>>>>>>>>>")
    print(doc.text())
    print("<<<<<<<<<<<<<<<<<<<<<<<<<<<")