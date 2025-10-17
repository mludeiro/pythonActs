from docreader import Document
from pathlib import Path
import docx

class DocXDocument(Document):
    def __init__(self, filepath: Path):
        super().__init__(filepath)
        self.document = docx.Document(self.filepath)

    @property
    def author(self) -> str:
        return self.document.core_properties.author

    def text(self) -> str:
        return "\n".join(para.text for para in self.document.paragraphs)